"""Build the Agent 01 user guide DOCX.

The companion PDF is built by build_agent01_user_guide_pdf.py.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "agent-01-blog-writing-agent-user-guide.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "E8EEF5"
PALE_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], LIGHT_FILL)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            set_cell_text(row.cells[idx], value)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    table.rows[0].cells[0].width = Inches(6.5)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, PALE_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLUE if level < 3 else DARK_BLUE
        run.font.name = "Calibri"


def add_para(doc: Document, text: str = "", *, style: str | None = None, bold_label: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    if bold_label:
        run = p.add_run(bold_label)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        add_para(doc, item, style="List Bullet")


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        add_para(doc, item, style="List Number")


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(code.strip("\n").splitlines()):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Agent 01 Blog Writing Agent Guide")


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Agent 01 Blog Writing Agent")
    run.font.size = Pt(24)
    run.font.color.rgb = DARK_BLUE
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("User Guide for the review-ready blog package UI")
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.add_run(f"Updated: {date.today().isoformat()}  |  Scope: Agent 01 v1  |  Mode: draft-only")

    add_callout(
        doc,
        "One-line summary",
        "Agent 01 turns messy text, voice recordings, or audio-only video input into a review-ready blog package. It drafts, reviews, scores, and flags issues, but it does not publish anywhere.",
    )

    add_heading(doc, "1. What Agent 01 Is Trying To Do")
    add_para(
        doc,
        "Agent 01 is the platform's golden/reference agent. It was built first so its architecture, provider abstractions, cost controls, tests, and UI pattern can be reused by later agents.",
    )
    add_para(
        doc,
        "For the user, its practical job is simple: take rough source material and produce a blog package that a human can review, edit lightly, and publish manually outside the agent.",
    )
    add_bullets(
        doc,
        [
            "Input: text, voice upload, video upload, or a structured Agent 03 blog brief.",
            "Process: normalize the source, extract ideas, plan the post, draft, review, and finalize.",
            "Output: title, summary, full draft, SEO keywords, tags, meta description, quality score, notes, and cost.",
            "Boundary: v1 is draft-only. It never publishes, posts to social media, writes to a CMS, scrapes the web, or analyzes video frames.",
        ],
    )

    add_heading(doc, "2. What It Can And Cannot Do")
    add_table(
        doc,
        ["Can do", "Cannot do in v1"],
        [
            ["Turn messy notes into a blog draft.", "Publish the blog or write to a CMS."],
            ["Transcribe voice recordings through the configured transcription provider.", "Use autonomous web search or live scraping."],
            ["Extract audio from video and transcribe it.", "Analyze video visuals, key frames, or screen content."],
            ["Handle pasted reference material as untrusted inspiration.", "Copy or spin source/reference material."],
            ["Track per-stage cost and stop before the Rs 50 ceiling.", "Guarantee factual claims without evidence."],
        ],
        [3.25, 3.25],
    )

    add_heading(doc, "3. How The Blog Pipeline Works")
    add_numbers(
        doc,
        [
            "Intake checks the input type and rejects unsupported or empty input before expensive work.",
            "For voice, the file is transcribed. For video, audio is extracted first, then transcribed. Video visuals are ignored in v1.",
            "Normalize cleans the content and separates usable source material from commands or weak input.",
            "Extract ideas identifies the main idea, supporting points, audience, and usable angle.",
            "Plan creates title candidates, outline, tone, audience, and keywords.",
            "Draft writes the blog using the plan and source material.",
            "Review scores the draft out of 100 and flags hard-fail issues.",
            "Finalize assembles the BlogPackage with status, draft, notes, quality, and cost.",
        ],
    )

    add_heading(doc, "4. Output Statuses")
    add_table(
        doc,
        ["Status", "Meaning", "What to do"],
        [
            ["pass", "Quality score is at least 80 and no hard-fail condition triggered.", "Review the draft, edit lightly, and publish manually outside the agent."],
            ["needs_human", "Input is too thin, unsafe, unsupported, or the review found a problem that needs a person.", "Add better source material or fix the issue noted in the result."],
            ["stopped_cost_ceiling", "The run stopped to protect the Rs 50 budget ceiling.", "Use shorter input or retry later with a cheaper path."],
            ["error", "A provider, config, file, or graph issue stopped the run.", "Check the terminal running uvicorn and verify GCP/env setup."],
        ],
        [1.35, 3.0, 2.15],
    )

    add_heading(doc, "5. How To Start The UI In Live GCP Mode")
    add_para(doc, "Open PowerShell and run these from the repository root:")
    add_code(
        doc,
        r'''
Set-Location "C:\Users\Pranesh\Desktop\agents-platform"

$repo = (Get-Location).Path
$env:BLOG_UI_PROVIDER = "gcp"
$env:VERTEX_AI_PROJECT = "agents-platform-1212"
$env:GCS_BLOG_BUCKET = "agents-platform-1212-agents-platform-stt-smoke"
$env:PYTHONPATH = "$repo\packages;$repo\agents\agent-01-blog-writer"

Set-Location "$repo\apps\blog-ui"
& "$repo\.agent02-ui-venv\Scripts\python.exe" -m uvicorn app:app --host 127.0.0.1 --port 8001
''',
    )
    add_para(doc, "Then open http://127.0.0.1:8001 in the browser.")
    add_callout(
        doc,
        "Important",
        "Live GCP mode makes billable Vertex AI and Speech-to-Text calls. Use mock mode only for developer checks; the normal user-facing test path should be GCP live.",
    )

    add_heading(doc, "6. How To Use The UI")
    add_numbers(
        doc,
        [
            "Choose the input type: Text, Voice upload, or Video upload.",
            "For Text, paste notes, transcript text, source material, or a structured brief.",
            "For Voice or Video, upload the file. The UI deletes the temporary upload after the graph returns.",
            "Optionally paste an Agent 03 blog brief JSON into the Agent 03 field.",
            "Click Generate Blog.",
            "Read the result page: status, notes, provider, total cost, quality score, stage costs, summary, and draft.",
        ],
    )

    add_heading(doc, "7. What Good Input Looks Like")
    add_para(doc, "Agent 01 performs best when the input includes the actual source idea, not only a command like 'write about AI'.")
    add_bullets(
        doc,
        [
            "Topic or working title.",
            "Target audience.",
            "Main idea or argument.",
            "3 to 6 supporting points.",
            "Tone and style preference.",
            "CTA or intended reader action.",
            "Any claims that need evidence, marked clearly as placeholders if not yet verified.",
        ],
    )
    add_heading(doc, "Copy-Paste Text Test Case", level=2)
    add_code(
        doc,
        """
Topic: How AI agents improve content marketing for B2B teams
Audience: B2B marketing managers and content leads
Tone: clear, practical, confident
Goal: explain a safe, human-reviewed content workflow

Source material:
B2B content teams often lose time turning scattered campaign notes, call
transcripts, and rough ideas into usable briefs. AI agents can help with
repeatable steps such as summarizing source material, creating outlines,
drafting first-pass sections, suggesting titles, and preparing review packages.
The best use case is not replacing editors. It is giving editors a cleaner
starting point so they can spend more time on judgment, accuracy, brand voice,
and final approval. Teams should start with one low-risk workflow, such as blog
draft preparation, and measure time saved before scaling.

CTA: Audit one repeatable content workflow this week.
""",
    )

    add_heading(doc, "8. Agent 03 Blog Brief Handoff")
    add_para(
        doc,
        "Agent 01 can accept a structured blog brief from Agent 03. When present, that brief becomes the primary source for audience, goal, angle, outline, CTA, keywords, constraints, evidence placeholders, and risk flags.",
    )
    add_code(
        doc,
        """
{
  "selected_idea_id": "idea_01",
  "selected_idea_title": "How AI Agents Help B2B Teams Scale Content",
  "suggested_title": "How AI Agents Help B2B Teams Scale Content",
  "title_options": [
    "How AI Agents Help B2B Teams Scale Content",
    "A Practical Guide to Human-Reviewed AI Content Workflows"
  ],
  "target_audience": "B2B marketing managers",
  "campaign_goal": "Build awareness for safe AI-assisted content workflows",
  "content_angle": "Show how agents reduce repetitive content operations while humans keep editorial control.",
  "core_message": "AI agents help content teams scale planning, drafting, and review without removing human judgment.",
  "pain_points": [
    "Campaign notes are scattered across calls, docs, and chats.",
    "Editors spend too much time preparing first drafts.",
    "Teams need speed without losing accuracy or brand voice."
  ],
  "value_proposition": "A repeatable agent-assisted workflow gives teams a cleaner first draft and a stronger review process.",
  "suggested_outline": [
    "Why content teams get stuck",
    "Where AI agents help",
    "Why human review still matters",
    "How to start safely"
  ],
  "proof_points_or_placeholders": [
    "Add internal time-saved evidence if available."
  ],
  "cta": "Audit one repeatable content workflow this week.",
  "tone": "clear, practical, confident",
  "keywords": ["AI agents", "content marketing", "content workflow"],
  "constraints": [
    "Do not invent statistics.",
    "Keep evidence placeholders visible when proof is missing."
  ],
  "risk_flags": ["evidence_placeholder_needed"],
  "quality_notes": ["Use specific workflow examples instead of hype."]
}
""",
    )

    add_heading(doc, "9. Cost, Quality, And Safety")
    add_bullets(
        doc,
        [
            "Cost ceiling: every blog run is protected by a hard Rs 50 ceiling.",
            "Quality pass: score must be at least 80 out of 100 and no hard-fail flag may be present.",
            "Revision loop: the graph may revise weak drafts within configured revision and cost caps.",
            "Prompt injection defense: pasted text, transcripts, and reference material are treated as untrusted data, never instructions.",
            "Originality: reference content can inspire structure or context, but the draft must not copy or spin it.",
            "Privacy: raw voice/video uploads are temporary and deleted after processing; run result JSON stays local under apps/blog-ui/runs.",
        ],
    )

    add_heading(doc, "10. Troubleshooting")
    add_table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["Page says GCP env vars missing", "The PowerShell window running uvicorn does not have the vars set.", "Stop uvicorn, set BLOG_UI_PROVIDER, VERTEX_AI_PROJECT, and GCS_BLOG_BUCKET, then restart."],
            ["Result says needs_human", "Input is too thin, too command-like, unsafe, or missing support.", "Add source material, supporting points, audience, tone, and CTA."],
            ["Provider is mock or output contains mock text", "Server started in mock/offline mode.", "Stop uvicorn, set BLOG_UI_PROVIDER=gcp, restart."],
            ["Voice or video fails", "Upload format, transcription, GCS, or GCP auth issue.", "Check terminal logs, verify ADC/GitHub/GCP setup, and retry with a small file."],
            ["Cost is high", "Long source material or strong-tier draft/review work.", "Use tighter input or split the task into smaller source chunks."],
        ],
        [1.75, 2.25, 2.5],
    )

    add_heading(doc, "11. Operator Checklist")
    add_bullets(
        doc,
        [
            "Use live GCP mode for real testing.",
            "Keep the browser form provider-free; provider mode belongs in environment variables.",
            "Do not commit run JSONs or uploaded media.",
            "Check that the final package status is pass before treating it as review-ready.",
            "Manually review facts, claims, brand voice, and citations before publishing outside the agent.",
        ],
    )

    add_heading(doc, "12. Where Files Live")
    add_table(
        doc,
        ["Area", "Path / behavior"],
        [
            ["Agent logic", "agents/agent-01-blog-writer/agent/"],
            ["UI wrapper", "apps/blog-ui/"],
            ["Run JSONs", "apps/blog-ui/runs/{run_id}.json"],
            ["Temporary uploads", "apps/blog-ui/uploads/ during a request, then deleted"],
            ["Shared providers", "packages/core/providers/"],
            ["Shared tests/evals", "packages/evals/ and agent tests"],
        ],
        [2.0, 4.5],
    )

    add_para(
        doc,
        "This guide is for operating Agent 01 v1. It intentionally keeps publishing, web search, scraping, vector retrieval, and visual video analysis outside scope.",
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
